"""
Document Processing Strategies
=============================

Strategy pattern implementation for document processing phase of the pipeline.
Different strategies can handle different types of document processing and 
create different Azure Cognitive Search index objects.

This module is specifically designed for the Personal Documentation Assistant MCP Server
and contains strategies tailored to work item-based document organization:
- PersonalDocumentationAssistantProcessingStrategy: Work item-focused processing
- Future strategies: Could include project-specific, category-specific, or other 
  domain-specific processing approaches for the Personal Documentation Assistant
"""

import json
import re
import asyncio
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional, Any, Iterator
from dataclasses import dataclass
import frontmatter
import os
from dotenv import load_dotenv

# Import our helper modules
import sys
sys.path.append(str(Path(__file__).parent.parent))

# File type specific parsing libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Load environment variables
load_dotenv()

# Azure Cognitive Search Index Fields - based on create_index.py schema
AZURE_SEARCH_INDEX_FIELDS = {
    'id', 'content', 'content_vector', 'file_path', 'file_name', 'file_type', 
    'title', 'tags', 'category', 'context_name', 'last_modified', 
    'chunk_index', 'metadata_json'
}


def _read_document_file(file_path: Path) -> Optional[Dict]:
    """
    Read and parse a document file with file-type specific parsing.
    Supports different file types with appropriate parsing libraries.
    
    Supports: .md (markdown), .txt (text), .docx (Word documents)
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Optional[Dict]: Dictionary with content and file_path, or None if failed
    """
    try:
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.docx':
            # Parse DOCX files using python-docx
            if not DOCX_AVAILABLE:
                print(f"Warning: python-docx not available, cannot parse {file_path}")
                return None
            
            content = _extract_docx_content(file_path)
            if not content:
                return None
                
        elif file_extension in ['.md', '.txt']:
            # Parse markdown and text files as UTF-8
            content = file_path.read_text(encoding='utf-8')
            
        else:
            # Unsupported file type
            print(f"Warning: Unsupported file type {file_extension} for {file_path}")
            return None

        if not content or not content.strip():
            return None  # Skip empty files

        return {
            'content': content.strip(),
            'file_path': str(file_path)
        }

    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return None


def _extract_docx_content(file_path: Path) -> Optional[str]:
    """
    Extract text content from a DOCX file with proper structure preservation.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Optional[str]: Extracted text content with heading structure preserved
    """
    try:
        doc = Document(file_path)
        content_parts = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Preserve heading structure by adding markdown-style formatting
            # This helps with title and heading extraction later
            if paragraph.style.name.startswith('Heading'):
                heading_level = paragraph.style.name.replace('Heading ', '').strip()
                try:
                    level = int(heading_level)
                    # Convert Word heading levels to markdown format
                    markdown_prefix = '#' * min(level, 6) + ' '
                    content_parts.append(f"{markdown_prefix}{text}")
                except ValueError:
                    # If heading level is not a number, treat as regular text
                    content_parts.append(text)
            else:
                content_parts.append(text)
        
        # Join with double newlines to preserve paragraph structure
        full_content = '\n\n'.join(content_parts)
        
        return full_content if full_content.strip() else None
        
    except Exception as e:
        print(f"Error extracting DOCX content from {file_path}: {e}")
        return None


def _simple_chunk_text(content: str, max_chunk_size: int = 4000, overlap: int = 200) -> List[str]:
    """
    Split text into chunks with simple sentence-based splitting.
    
    Args:
        content: Text content to chunk
        max_chunk_size: Maximum characters per chunk
        overlap: Number of overlapping characters between chunks
        
    Returns:
        List[str]: List of text chunks
    """
    if not content or len(content) <= max_chunk_size:
        return [content] if content else []
    
    chunks = []
    sentences = content.split('. ')
    
    current_chunk = ""
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
            
        # Add period back if it was removed by split (except for last sentence)
        if not sentence.endswith('.') and not sentence.endswith('!') and not sentence.endswith('?'):
            sentence += '.'
        
        # Check if adding this sentence would exceed the limit
        test_chunk = current_chunk + " " + sentence if current_chunk else sentence
        
        if len(test_chunk) <= max_chunk_size:
            current_chunk = test_chunk
        else:
            # Save current chunk and start new one
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence
    
    # Add the last chunk
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    # If we still have chunks that are too long, split them more aggressively
    final_chunks = []
    for chunk in chunks:
        if len(chunk) <= max_chunk_size:
            final_chunks.append(chunk)
        else:
            # Split by words if sentences are too long
            words = chunk.split()
            current_word_chunk = ""
            for word in words:
                test_chunk = current_word_chunk + " " + word if current_word_chunk else word
                if len(test_chunk) <= max_chunk_size:
                    current_word_chunk = test_chunk
                else:
                    if current_word_chunk:
                        final_chunks.append(current_word_chunk.strip())
                    current_word_chunk = word
            if current_word_chunk:
                final_chunks.append(current_word_chunk.strip())
    
    return final_chunks


@dataclass
class ProcessedDocument:
    """A document that has been processed and is ready for search index upload."""
    # Core identification
    document_id: str
    file_path: str
    file_name: str
    file_type: str
    
    # Content
    title: str
    content: str
    content_chunks: List[str]
    
    # Metadata
    tags: List[str]
    category: Optional[str]
    context_name: Optional[str]
    
    # Timestamps
    last_modified: str
    
    # Processing metadata
    chunk_count: int
    processing_strategy: str
    metadata_json: Optional[str]  # Additional strategy-specific metadata
    
    @property
    def metadata(self) -> Dict[str, Any]:
        """
        Return structured metadata for this processed document.
        This property provides a consolidated view of all document metadata
        for use with file tracking and audit systems.
        
        Returns:
            Dict containing document metadata including:
            - Basic file information (title, file_type, file_name)
            - Categorization (category, context_name, tags)
            - Processing details (processing_strategy, chunk_count)
            - Timestamps and additional metadata
        """
        metadata = {
            'title': self.title,
            'file_type': self.file_type,
            'file_name': self.file_name,
            'category': self.category,
            'context_name': self.context_name,
            'tags': self.tags if self.tags else [],
            'processing_strategy': self.processing_strategy,
            'chunk_count': self.chunk_count,
            'last_modified': self.last_modified,
            'document_id': self.document_id,
        }
        
        # Include additional strategy-specific metadata if available
        if self.metadata_json:
            try:
                import json
                additional_metadata = json.loads(self.metadata_json)
                metadata.update(additional_metadata)
            except (json.JSONDecodeError, TypeError):
                # If JSON parsing fails, include raw metadata_json
                metadata['metadata_json'] = self.metadata_json
        
        return metadata


@dataclass
class DocumentProcessingResult:
    """Result of the document processing phase."""
    total_documents: int
    successfully_processed: int
    failed_documents: int
    processed_documents: List[ProcessedDocument]
    processing_time: float
    errors: List[str]
    strategy_name: str
    strategy_metadata: Optional[Dict[str, Any]] = None


class DocumentProcessingStrategy(ABC):
    """
    Abstract base class for document processing strategies.
    
    Each strategy defines how to:
    1. Process discovered files into ProcessedDocument objects
    2. Generate Azure Cognitive Search index objects
    3. Handle strategy-specific metadata and indexing requirements
    """
    
    @abstractmethod
    def get_strategy_name(self) -> str:
        """
        Get the name of this processing strategy.
        
        Returns:
            str: Strategy name for identification and logging
        """
        pass
    
    @abstractmethod
    def process_documents(self, discovered_files: List[Path]) -> DocumentProcessingResult:
        """
        Process a list of discovered files into ProcessedDocument objects.
        
        Args:
            discovered_files: List of file paths from discovery phase
            
        Returns:
            DocumentProcessingResult: Processing results with processed documents
        """
        pass
    
    @abstractmethod
    def create_search_index_objects(self, processed_documents: List[ProcessedDocument]) -> Iterator[Dict[str, Any]]:
        """
        Create Azure Cognitive Search index objects from processed documents.
        
        Different strategies may create different index schemas or object structures
        based on their specific requirements.
        
        Args:
            processed_documents: List of processed documents
            
        Yields:
            Dict[str, Any]: Search index object for Azure Cognitive Search
        """
        pass
    
    def process_single_document(self, file_path: Path) -> Optional[ProcessedDocument]:
        """
        Process a single document file. This method can be overridden by subclasses
        for strategy-specific processing, but provides a default implementation.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ProcessedDocument: Processed document ready for search upload, or None if failed
        """
        # Read the document
        file_data = _read_document_file(file_path)
        if not file_data:
            return None
        
        content = file_data['content']
        
        # Extract metadata using strategy-specific logic
        metadata = self.extract_metadata(content, file_path)
        
        # Generate chunks using strategy-specific chunking
        chunks = self.chunk_document_content(content) if hasattr(self, 'chunk_document_content') else _simple_chunk_text(content)
        
        # Map metadata to search index fields
        context_name = self.extract_context_info(metadata)
        
        # Prepare additional metadata as JSON
        additional_metadata = self.prepare_additional_metadata(metadata)
        
        return ProcessedDocument(
            document_id=self.generate_document_id(file_path),
            file_path=str(file_path),
            file_name=file_path.name,
            file_type=metadata.get('file_type', 'unknown'),
            title=metadata.get('title', file_path.stem),
            content=content,
            content_chunks=chunks,
            tags=metadata.get('tags', []),
            category=metadata.get('category', metadata.get('document_category')),
            context_name=context_name,
            last_modified=metadata.get('last_modified', datetime.now().isoformat() + 'Z'),
            chunk_count=len(chunks),
            processing_strategy=self.get_strategy_name(),
            metadata_json=json.dumps(additional_metadata) if additional_metadata else None
        )
    
    @abstractmethod
    def extract_metadata(self, content: str, file_path: Path) -> Dict:
        """
        Extract comprehensive metadata from file content and path.
        This is strategy-specific and implements the metadata extraction approach
        outlined in EXTRACTING_FILE_METADATA.md.
        
        Args:
            content: Raw file content
            file_path: Path to the file
            
        Returns:
            Dict: Extracted metadata including title, tags, work_item_id, etc.
        """
        pass
    
    @abstractmethod
    def extract_context_info(self, metadata: Dict) -> Optional[str]:
        """
        Extract context name from metadata based on the strategy.
        
        Args:
            metadata: Document metadata
            
        Returns:
            Optional[str]: Context name (work items, projects, etc.)
        """
        pass
    
    def prepare_additional_metadata(self, metadata: Dict) -> Dict:
        """
        Prepare additional metadata that doesn't fit in standard search fields.
        Can be overridden by subclasses for strategy-specific metadata handling.
        
        Args:
            metadata: Full document metadata
            
        Returns:
            Dict: Additional metadata to store as JSON
        """
        # Everything else goes to additional metadata
        additional = {k: v for k, v in metadata.items() if k not in AZURE_SEARCH_INDEX_FIELDS}
        
        return additional if additional else {}
    
    def generate_document_id(self, file_path: Path) -> str:
        """
        Generate a unique document ID. Can be overridden by subclasses.
        
        Args:
            file_path: Path to the document
            
        Returns:
            str: Unique document ID
        """
        # Use file path hash for consistency
        import hashlib
        path_str = str(file_path).replace('\\', '/')
        return hashlib.md5(path_str.encode()).hexdigest()


class PersonalDocumentationAssistantProcessingStrategy(DocumentProcessingStrategy):
    """
    Processing strategy for Personal Documentation Assistant use case.
    
    Focuses on work item organization and creates index objects optimized 
    for work item-based document querying and retrieval.
    """
    
    def get_strategy_name(self) -> str:
        return "PersonalDocumentationAssistant"
    
    def process_documents(self, discovered_files: List[Path]) -> DocumentProcessingResult:
        """Process documents with Personal Documentation Assistant focus."""
        start_time = datetime.now()
        processed_documents = []
        errors = []
        successfully_processed = 0
        failed_documents = 0
        
        # Track work item statistics
        work_items_found = set()
        
        print(f"      🔄 Processing {len(discovered_files)} discovered documents individually...")
        
        for i, file_path in enumerate(discovered_files, 1):
            try:
                print(f"         {i:2d}/{len(discovered_files)}: Processing {file_path.name}...", end=" ")
                
                processed_doc = self.process_single_document(file_path)
                if processed_doc:
                    processed_documents.append(processed_doc)
                    successfully_processed += 1
                    
                    # Track work items
                    if processed_doc.context_name:
                        work_items_found.add(processed_doc.context_name)
                    
                    print(f"✅ ({processed_doc.chunk_count} chunks)")
                else:
                    failed_documents += 1
                    errors.append(f"Failed to process: {file_path}")
                    print("❌ Failed to process")
                    
            except Exception as e:
                failed_documents += 1
                errors.append(f"Error processing {file_path}: {str(e)}")
                print(f"❌ Error: {str(e)}")
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        # Strategy-specific metadata
        strategy_metadata = {
            "work_items_count": len(work_items_found),
            "work_items_found": list(work_items_found),
            "documents_per_work_item": {
                work_item: len([doc for doc in processed_documents if doc.context_name == work_item])
                for work_item in work_items_found
            }
        }
        
        return DocumentProcessingResult(
            total_documents=len(discovered_files),
            successfully_processed=successfully_processed,
            failed_documents=failed_documents,
            processed_documents=processed_documents,
            processing_time=processing_time,
            errors=errors,
            strategy_name=self.get_strategy_name(),
            strategy_metadata=strategy_metadata
        )
    
    def extract_context_info(self, metadata: Dict) -> Optional[str]:
        """Extract work item information as context."""
        context_name = metadata.get('work_item_id')  # Could be enhanced with actual work item names
        return context_name
    
    def extract_metadata(self, content: str, file_path: Path) -> Dict:
        """
        Extract comprehensive metadata using Personal Documentation Assistant strategy.
        
        Implements the strategy outlined in EXTRACTING_FILE_METADATA.md:
        1. Directory name provides context (work item ID) - becomes keyword for all files
        2. Headings in markdown/text files become keywords
        3. File type determines extraction approach
        4. Directory structure provides hierarchical context
        """
        try:
            # Determine file type
            file_extension = file_path.suffix.lower()
            file_type = self._get_file_type(file_extension)
            
            # Initialize metadata
            metadata = {}
            
            # Extract frontmatter if markdown file
            post = None
            clean_content = content
            if file_type == 'markdown':
                try:
                    post = frontmatter.loads(content)
                    metadata = dict(post.metadata) if post.metadata else {}
                    clean_content = post.content if post.metadata else content
                except:
                    pass  # If frontmatter parsing fails, use content as-is
            
            # Extract work item ID from directory structure (primary strategy)
            work_item_id = self._extract_work_item_id(file_path)
            metadata['work_item_id'] = work_item_id
            metadata['file_type'] = file_type
            
            # Add the cleaned content to metadata
            metadata['content'] = clean_content
            
            # Extract title using strategy priority
            metadata['title'] = self._extract_title(clean_content, file_path, metadata, file_type)
            
            # Extract tags/keywords using comprehensive strategy
            metadata['tags'] = self._extract_tags(clean_content, file_path, metadata, file_type, work_item_id, post)
            
            # Add file system metadata
            metadata.update(self._extract_file_system_metadata(file_path))
            
            return metadata
            
        except Exception as e:
            # Return minimal metadata on error
            return self._get_error_metadata(file_path, str(e))
    
    def _get_file_type(self, file_extension: str) -> str:
        """Map file extension to file type."""
        type_mapping = {
            '.md': 'markdown',
            '.txt': 'text',
            '.docx': 'document'
        }
        return type_mapping.get(file_extension, 'unknown')
    
    def _extract_work_item_id(self, file_path: Path) -> str:
        """
        Extract work item ID from directory structure.
        
        Strategy from EXTRACTING_FILE_METADATA.md: 
        "Directory name provides valuable information about the context of files.
        For example if directory name -> 'Task 12', it means files under this are 
        related to 'Task 12' and 'Task 12' becomes a keyword of interest."
        """
        # Use parent directory name as work item ID (primary context)
        work_item_dir = file_path.parent
        return work_item_dir.name
    
    def _extract_title(self, content: str, file_path: Path, metadata: Dict, file_type: str) -> str:
        """
        Extract title using Personal Documentation Assistant strategy.
        
        Current implementation: Use filename for simplicity and consistency.
        
        Future enhancement: Scan document content using LLM to identify correct titles
        by analyzing headings, document structure, and semantic content.
        """
        # Use filename with better formatting as title (current implementation)
        filename_title = file_path.stem.replace('_', ' ').replace('-', ' ')
        formatted_title = ' '.join(word.capitalize() for word in filename_title.split())
        
        # TODO: Future enhancement - LLM-based title extraction
        # Plan: Pass document content through LLM (OpenAI/Azure OpenAI) to:
        # 1. Analyze document structure and headings
        # 2. Identify the most appropriate title based on content
        # 3. Handle different document types (markdown, DOCX, text) intelligently
        # 4. Extract semantic meaning to generate descriptive titles
        # Example implementation:
        # llm_title = self._extract_llm_title(content, file_type, filename_title)
        # return llm_title if llm_title else formatted_title
        
        return formatted_title
    
    def _extract_tags(self, content: str, file_path: Path, metadata: Dict, 
                     file_type: str, work_item_id: str, post) -> list:
        """
        Extract tags/keywords using Personal Documentation Assistant strategy.
        
        Current implementation: Basic tags with work item ID, file type, and filename.
        
        Future enhancement: Pass document through LLM to extract intelligent keywords
        and add them to the tags list for better searchability and categorization.
        """
        # Initialize basic tags
        tags = set()
        
        # Add work item ID as primary tag
        tags.add(work_item_id.lower())
        
        # Add file type as tag
        tags.add(file_type)
        
        # Add filename (stem without extension) as tag for better searchability
        filename_tag = file_path.stem.lower().replace('_', '-').replace(' ', '-')
        tags.add(filename_tag)
        
        # TODO: Future enhancement - LLM-based keyword extraction
        # Plan: Pass document content through LLM (OpenAI/Azure OpenAI) to extract
        # intelligent keywords, technical terms, and domain-specific concepts
        # Example implementation:
        # llm_keywords = self._extract_llm_keywords(content, file_type)
        # tags.update(llm_keywords)
        
        # Return sorted tags (currently work_item_id, file_type, and filename)
        return sorted(list(tags))
    
    def _extract_file_system_metadata(self, file_path: Path) -> Dict:
        """Extract file system metadata."""
        file_stat = file_path.stat()
        last_modified_dt = datetime.fromtimestamp(file_stat.st_mtime)
        
        return {
            'last_modified': last_modified_dt.isoformat() + 'Z',
            'work_item_directory': str(file_path.parent),
            'directory_path': str(file_path.parent)
        }
       
    def _get_error_metadata(self, file_path: Path, error_message: str) -> Dict:
        """Return minimal metadata when extraction fails."""
        # Create filename tag for consistency with successful processing
        filename_tag = file_path.stem.lower().replace('_', '-').replace(' ', '-')
        
        return {
            'file_type': 'unknown',
            'title': file_path.stem,
            'tags': [file_path.parent.name, 'error', filename_tag],
            'work_item_id': file_path.parent.name,
            'last_modified': datetime.now().isoformat() + 'Z',
            'extraction_error': error_message
        }
    
    def chunk_document_content(self, content: str, max_chunk_size: int = 4000) -> List[str]:
        """
        Chunk document content into manageable pieces for embedding generation.
        
        This method provides strategy-specific chunking that can be optimized
        for the Personal Documentation Assistant use case.
        
        Args:
            content: Full document content to chunk
            max_chunk_size: Maximum size per chunk in characters
            
        Returns:
            List[str]: List of content chunks ready for embedding generation
        """
        return _simple_chunk_text(content, max_chunk_size)
    
    async def generate_chunk_embeddings(self, chunks: List[str]) -> List[List[float]]:
        """
        Generate embeddings for document chunks using Azure OpenAI.
        
        This method handles the embedding generation process and provides
        fallback handling for failed embeddings.
        
        Args:
            chunks: List of text chunks to generate embeddings for
            
        Returns:
            List[List[float]]: List of embedding vectors, empty list for failed embeddings
        """
        from common.embedding_service import get_embedding_generator
        
        embedding_generator = get_embedding_generator()
        embeddings = []
        
        for i, chunk in enumerate(chunks):
            try:
                embedding = await embedding_generator.generate_embedding(chunk)
                if embedding is None:
                    print(f"Warning: Failed to generate embedding for chunk {i}")
                    embeddings.append([])  # Empty vector as fallback
                else:
                    embeddings.append(embedding)
            except Exception as e:
                print(f"Error generating embedding for chunk {i}: {e}")
                embeddings.append([])  # Empty vector as fallback
        
        return embeddings
    
    def create_chunk_search_objects(self, processed_doc: ProcessedDocument, 
                                   chunk_embeddings: List[List[float]]) -> List[Dict[str, Any]]:
        """
        Create Azure Cognitive Search index objects for each chunk of a document.
        
        Combines document metadata with chunk content and embeddings to create
        search index objects that are ready for Azure Cognitive Search upload.
        
        Args:
            processed_doc: Processed document with metadata and chunks
            chunk_embeddings: List of embedding vectors for each chunk
            
        Returns:
            List[Dict[str, Any]]: List of search index objects, one per chunk
        """
        search_objects = []
        
        for chunk_index, (chunk_content, embedding) in enumerate(zip(processed_doc.content_chunks, chunk_embeddings)):
            # Convert tags array to comma-separated string (Azure Search expects string, not array)
            tags_str = ', '.join(processed_doc.tags) if isinstance(processed_doc.tags, list) else str(processed_doc.tags) if processed_doc.tags else ''
            
            # Create enhanced chunk index with file path and chunk number for better identification
            # Format: "filename.ext_chunk_N" for easy identification and sorting
            enhanced_chunk_index = f"{processed_doc.file_name}_chunk_{chunk_index}"
            
            # Create search index object matching the exact schema from create_index.py
            search_object = {
                "id": f"{processed_doc.document_id}_chunk_{chunk_index}",
                "content": chunk_content,
                "content_vector": embedding,
                "file_path": processed_doc.file_path,
                "file_name": processed_doc.file_name,
                "file_type": processed_doc.file_type,
                "title": processed_doc.title,
                "tags": tags_str,  # Convert to string format
                "category": processed_doc.category,
                "context_name": processed_doc.context_name,
                "last_modified": processed_doc.last_modified,
                "chunk_index": enhanced_chunk_index,
                "metadata_json": processed_doc.metadata_json
            }
            search_objects.append(search_object)
        
        return search_objects
    
    async def create_search_index_objects(self, processed_documents: List[ProcessedDocument]):
        """
        Create Azure Cognitive Search index objects optimized for Personal Documentation Assistant.
        
        Creates chunk-based objects with embeddings, work item context, and enhanced metadata.
        Each chunk gets its own embedding for optimal vector search performance.
        """
        for doc in processed_documents:
            # Generate embeddings for all chunks of this document
            try:
                chunk_embeddings = await self.generate_chunk_embeddings(doc.content_chunks)
            except Exception as e:
                print(f"Error generating embeddings for document {doc.file_name}: {e}")
                # Create empty embeddings as fallback
                chunk_embeddings = [[] for _ in doc.content_chunks]
            
            # Create search objects for each chunk with its embedding and metadata
            search_objects = self.create_chunk_search_objects(doc, chunk_embeddings)
            
            # IMPORTANT: Using yield here implements a generator pattern for memory efficiency
            # 
            # Why yield is crucial for this pipeline:
            # 1. MEMORY EFFICIENCY: Instead of creating all search objects in memory at once
            #    (which could be 100+ documents × 3-5 objects each = 300-500+ objects),
            #    yield returns objects one at a time, keeping memory usage minimal
            #
            # 2. STREAMING PROCESSING: The upload phase can process and upload objects
            #    immediately as they're generated, rather than waiting for all objects
            #    to be created first. This enables true streaming document processing.
            #
            # 3. SCALABILITY: For large document collections, this prevents memory overflow
            #    and allows the pipeline to handle thousands of documents efficiently
            #
            # 4. EARLY TERMINATION: If an error occurs during upload, processing can
            #    stop immediately without wasting resources on remaining objects
            #
            # Example flow:
            # for search_object in create_search_index_objects(docs):  # One at a time
            #     upload_to_azure(search_object)  # Process immediately
            #     # Only 1 object in memory at any time vs 300+ without yield

            # Yield each search object
            for search_object in search_objects:
                yield search_object
